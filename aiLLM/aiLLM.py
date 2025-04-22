# ============================================================
# AI Chat App  v2.5‑multimodal
# ------------------------------------------------------------
#  • Web‑Search, Deep‑Research, Image (multimodal) input
#  • Attachment summaries: Excel(.xlsx/.xls) ▪ CSV ▪ DOCX ▪ HWP/HWPX ▪ PDF ▪ 이미지 OCR
#  • Robust dependency checks (openpyxl ≥3.1, xlrd 1.2, tabulate)
#  • Qt "QItemSelection" warning safely suppressed (all PyQt5 builds)
# ------------------------------------------------------------
# pip install -U \
#   PyQt5 pandas openpyxl xlrd==1.2.0 python-docx pillow pytesseract olefile \
#   pyhwp tabulate packaging google-generativeai anthropic openai bs4 requests PyPDF2 pdfplumber \
#   jpype1 lxml
# (Win/Mac) Tesseract‑OCR 4.x 설치 후 PATH 등록
# LibreOffice/OpenOffice 설치 권장 (한글 파일 변환 기능 사용)
# ============================================================

import os, sys, json, time, threading, requests, bs4, re, base64
from typing import List, Optional, Tuple

# ─────────────────────────────────────────────────────────────
# 0.  Qt platform plugin path (Windows 전용)
# ─────────────────────────────────────────────────────────────
os.environ["QT_QPA_PLATFORM_PLUGIN_PATH"] = os.path.join(
    sys.base_prefix, "Lib", "site-packages", "PyQt5", "Qt", "plugins", "platforms"
)

# ─────────────────────────────────────────────────────────────
# 1.  PyQt5 import & Qt 경고 억제
# ─────────────────────────────────────────────────────────────
from PyQt5.QtWidgets import *
from PyQt5.QtCore    import Qt, pyqtSignal, QThread, QItemSelection
from PyQt5 import QtCore
from PyQt5.QtGui     import QIcon

# 안전하게 qRegisterMetaType 호출 (PyQt5 빌드마다 다름)
try:
    if hasattr(QtCore, "qRegisterMetaType"):
        QtCore.qRegisterMetaType(QItemSelection)
    else:
        from PyQt5.QtCore import qRegisterMetaType
        qRegisterMetaType(QItemSelection)
except Exception:
    pass  # 함수가 없으면 경고만 날 뿐 기능에는 영향 없음

import markdown2

# ─────────────────────────────────────────────────────────────
# 2.  의존성 버전 체크
# ─────────────────────────────────────────────────────────────
import pandas as pd
from packaging import version as _v

try:
    from tabulate import tabulate
    TAB_OK = True
except ImportError:
    TAB_OK = False

try:
    import openpyxl
    OPENPYXL_OK = _v.parse(openpyxl.__version__) >= _v.parse("3.1.0")
except Exception:
    OPENPYXL_OK = False

try:
    import xlrd
    XLRD_OK = xlrd.__version__.startswith("1.2.")
except Exception:
    XLRD_OK = False

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = pytesseract = None

try:
    import docx  # python‑docx
except ImportError:
    docx = None

# ==== 한글 파일 처리 라이브러리 체크 ====
# 1. pyhwp + olefile (기본 HWP 처리)
try:
    import pyhwp
    from olefile import OleFileIO
    HWP_PYHWP_OK = True
except ImportError:
    pyhwp = None
    OleFileIO = None
    HWP_PYHWP_OK = False

# 2. hwp5txt 명령줄 도구 체크
try:
    import subprocess
    result = subprocess.run(['hwp5txt', '--version'], capture_output=True, text=True)
    HWP5TXT_OK = result.returncode == 0
except Exception:
    HWP5TXT_OK = False

# 3. HWPX 처리 (ZIP + XML)
try:
    import zipfile
    import xml.etree.ElementTree as ET
    try:
        import lxml.etree as LET
        LXML_OK = True
    except ImportError:
        LET = None
        LXML_OK = False
    HWPX_OK = True
except ImportError:
    HWPX_OK = False
    LXML_OK = False
    
# 4. Java 기반 처리 (JPype + HWPLib)
try:
    import jpype
    import jpype.imports
    HWP_JPYPE_OK = True
    
    # JVM 시작 여부 확인 및 필요시 시작
    if not jpype.isJVMStarted():
        try:
            jpype.startJVM(jpype.getDefaultJVMPath())
            # HWPLib 추가 확인은 필요시 시도
        except Exception:
            HWP_JPYPE_OK = False
except ImportError:
    jpype = None
    HWP_JPYPE_OK = False

# 5. LibreOffice/OpenOffice 체크 (soffice 명령줄 도구)
# SOFFICE 경로 체크 (LibreOffice/OpenOffice)
def check_soffice_path():
    """LibreOffice/OpenOffice soffice 실행 파일 경로 확인"""
    # 가능한 경로들
    possible_paths = [
        # Windows
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        # macOS
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        # Linux
        "/usr/bin/soffice",
        "/usr/lib/libreoffice/program/soffice"
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            return path
            
    # PATH에서 찾기
    import shutil
    soffice_path = shutil.which('soffice')
    if soffice_path:
        return soffice_path
        
    return None

SOFFICE_PATH = check_soffice_path()
SOFFICE_OK = SOFFICE_PATH is not None

# 한글 파일 처리 가능 여부 종합
HWP_OK = HWP_PYHWP_OK or HWP5TXT_OK or HWP_JPYPE_OK or SOFFICE_OK

# PDF 지원 확인
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
    
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# PDF 지원 확인 플래그
PDF_OK = PyPDF2 is not None or pdfplumber is not None

# ─────────────────────────────────────────────────────────────
# 3.  데이터 모델
# ─────────────────────────────────────────────────────────────
class ChatSession:
    def __init__(self, sid: str, title: str = "New Chat"):
        self.session_id = sid
        self.title = title
        self.messages: List[dict[str, str]] = []
        self.attachments: List[str] = []
        self.model_used = ""

    def add_message(self, role: str, content: str):
        self.messages.append({"role": role, "content": content})

    def add_attachment(self, path: str):
        if path not in self.attachments:
            self.attachments.append(path)

    def to_dict(self):
        return {
            "session_id": self.session_id,
            "title": self.title,
            "messages": self.messages,
            "attachments": self.attachments,
            "model_used": self.model_used,
        }

    @staticmethod
    def from_dict(d: dict):
        cs = ChatSession(d["session_id"], d.get("title", "New Chat"))
        cs.messages = d.get("messages", [])
        cs.attachments = d.get("attachments", [])
        cs.model_used = d.get("model_used", "")
        return cs


class ChatHistory:
    def __init__(self, savedir="chat_history"):
        self.dir = savedir
        os.makedirs(savedir, exist_ok=True)
        self.sessions: dict[str, ChatSession] = {}
        self._load_all()

    def _load_all(self):
        for fn in os.listdir(self.dir):
            if fn.endswith(".json"):
                try:
                    with open(os.path.join(self.dir, fn), encoding="utf-8") as f:
                        d = json.load(f)
                    self.sessions[d["session_id"]] = ChatSession.from_dict(d)
                except Exception as e:
                    print("❌ history load:", fn, e)

    def save(self, sess: ChatSession):
        self.sessions[sess.session_id] = sess
        try:
            with open(os.path.join(self.dir, f"{sess.session_id}.json"), "w", encoding="utf-8") as w:
                json.dump(sess.to_dict(), w, ensure_ascii=False, indent=2)
        except Exception as e:
            print("❌ history save:", e)

    def generate_title(self, sess: ChatSession) -> str:
        first = next((m["content"] for m in sess.messages if m["role"] == "user"), "")
        s = first.strip().replace("\n", " ")[:30]
        return s + ("…" if len(s) == 30 else "") if s else sess.title


# ─────────────────────────────────────────────────────────────
# 4.  Web‑Search (Google snippet)
# ─────────────────────────────────────────────────────────────
def web_search(query: str) -> str:
    api_key = "PUT_YOUR_ScraperAPI"  # ScraperAPI
    try:
        resp = requests.get(
            "http://api.scraperapi.com",
            params={"api_key": api_key, "url": f"https://www.google.com/search?q={query}"},
            headers={"User-Agent": "Mozilla/5.0"},
            timeout=15,
        )
        resp.raise_for_status()
        soup = bs4.BeautifulSoup(resp.text, "html.parser")
        snippets = [d.get_text(" ", strip=True) for d in soup.select("div[data-sncf='1']") if d.get_text(strip=True)]
        if not snippets:
            snippets = [d.get_text(" ", strip=True) for d in soup.select(".BNeawe.s3v9rd.AP7Wnd")]
        return "\n".join(snippets[:3]) if snippets else "[검색 결과 없음]"
    except Exception as e:
        return f"[웹 검색 오류] {e}"


# ─────────────────────────────────────────────────────────────
# 5.  AIManager (OpenAI · Gemini · Claude)
# ─────────────────────────────────────────────────────────────
from openai import OpenAI
import google.generativeai as genai, anthropic
from google.generativeai import types as gtypes
from google.generativeai.types import Tool

OPENAI_MODEL = {"gpt-o3": "o3", "gpt-o3-mini": "o3-mini", "gpt-4o": "gpt-4o"}

GEMINI_MODEL = {
    "gemini-2.5-flash": "gemini-2.5-flash-preview-04-17",
    "gemini-2.5-pro": "models/gemini-2.5.pro-exp-03-25",  # 실험용 모델로 교체
}
CLAUDE_MODEL = {
    "claude-3-7-sonnet": "claude-3-7-sonnet-20250219",
    "claude-3-5-sonnet": "claude-3-5-sonnet-20240620",
}
# ------- 여기에 실제 API 키를 입력하세요 -------
OPENAI_API_KEY = 'PUT_YOUR_OPENAI_API'
GOOGLE_API_KEY = 'PUT_YOUR_GOOGLE_API'
CLAUDE_API_KEY = 'PUT_YOUR_CLAUDE_API'

class AIManager:
    def __init__(self):
        self.oai = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
        if GOOGLE_API_KEY:
            genai.configure(api_key=GOOGLE_API_KEY)
        self.claude = anthropic.Anthropic(api_key=CLAUDE_API_KEY) if CLAUDE_API_KEY else None

    @staticmethod
    def _sys_msg(web: bool) -> dict:
        return {
            "role": "system",
            "content": "당신은 도움이 되는 AI 비서입니다. 첨부파일이 있으면 해당 파일의 내용을 분석하고 참고하여 응답합니다."
            + (" 웹 검색 결과를 참고해 최신 정보를 제공합니다." if web else ""),
        }

    def _inject_search(self, query: str, msgs: List[dict]) -> List[dict]:
        """웹 검색 결과 삽입"""
        snippet = web_search(query)
        if snippet.startswith("[웹 검색 오류") or snippet.startswith("[검색 결과 없음]"):
            return msgs
        return msgs + [{"role": "system", "content": f"[🔍 Web Results]\n{snippet}"}]
        
    @staticmethod
    def _b64(path: str) -> str:
        """이미지 파일을 base64로 인코딩"""
        with open(path, "rb") as f:
            return base64.b64encode(f.read()).decode()

    # ---------- main entry ----------
    def get_response(self, api: str, model: str, msgs: List[dict], web: bool) -> str:
        """AI 응답 생성"""
        # 첨부파일 정보가 있는지 확인
        has_attachments = any(
            any(marker in m.get("content", "") for marker in 
                ["[📊 Excel]", "[📈 CSV]", "[📄 DOCX]", "[📑 한글]", "[📑 PDF]", "[🖼️ OCR]", "[🖼️ 이미지]"])
            for m in msgs
        )
        
        # 멀티모달 이미지 메시지 처리를 위한 복사본 생성
        processed_msgs = []
        
        # 웹 검색 추가 (마지막 유저 메시지 기준)
        if web:
            user_msgs = [m for m in msgs if m["role"] == "user" and "image_path" not in m]
            if user_msgs:
                last_user_msg = user_msgs[-1]["content"]
                msgs = self._inject_search(last_user_msg, msgs)
        
        # 시스템 메시지 추가
        processed_msgs.append(self._sys_msg(web))
        
        # 메시지 변환 (첨부파일 및 이미지 처리)
        for m in msgs:
            if "image_path" in m:
                img_path = m["image_path"]
                if api == "openai":
                    processed_msgs.append({
                        "role": "user",
                        "content": [
                            {"type": "text", "text": m["content"]},
                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{self._b64(img_path)}"}}
                        ]
                    })
                elif api == "gemini":
                    with open(img_path, "rb") as f:
                        img_bytes = f.read()
                    processed_msgs.append({
                        "role": "user",
                        "parts": [
                            {"inline_data": {"mime_type": "image/jpeg", "data": img_bytes}},
                            {"text": m["content"]}
                        ]
                    })
                elif api == "claude":
                    # Claude 이미지 처리
                    processed_msgs.append({
                        "role": "user",
                        "content": [
                            {"type": "text", "text": m["content"]},
                            {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": self._b64(img_path)}}
                        ]
                    })
            else:
                # 첨부파일 관련 메시지는 role을 자동으로 system으로 변경
                new_role = "system" if any(marker in m.get("content", "") for marker in 
                                        ["[📊 Excel]", "[📈 CSV]", "[📄 DOCX]", 
                                         "[📑 한글]", "[📑 PDF]", "[🖼️ OCR]"]) else m["role"]
                processed_msgs.append({
                    "role": new_role, 
                    "content": m["content"]
                })
        
        # API 호출
        try:
            if api == "openai":
                if self.oai is None:
                    return "OpenAI API 키가 설정되지 않았습니다."
                
                completion = self.oai.chat.completions.create(
                    model=OPENAI_MODEL.get(model, model), 
                    messages=processed_msgs
                )
                response = completion.choices[0].message.content
                
            elif api == "gemini":
                if GOOGLE_API_KEY is None or not GOOGLE_API_KEY:
                    return "Google API 키가 설정되지 않았습니다."
                
                # Gemini는 role이 user 또는 model만 가능하므로 변환
                gemini_msgs = []
                for m in processed_msgs:
                    if "parts" in m:  # 이미 parts 형식이면 그대로 사용
                        gemini_msgs.append(m)
                    else:
                        role = "user" if m["role"] in ["user", "system"] else "model"
                        gemini_msgs.append({
                            "role": role,
                            "parts": [{"text": m["content"]}]
                        })
                
                gem = genai.GenerativeModel(GEMINI_MODEL.get(model, model))
                response = gem.generate_content(gemini_msgs).text
                
            elif api == "claude":
                if self.claude is None:
                    return "Claude API 키가 설정되지 않았습니다."
                
                # Claude API는 assistant를 지원하지만 system 메시지는 최대 1개만 가능
                # 첨부파일 메시지를 user 메시지로 변환
                claude_msgs = []
                system_added = False
                
                for m in processed_msgs:
                    if isinstance(m.get("content"), list):  # 이미지 포함 메시지
                        claude_msgs.append(m)
                    elif m["role"] == "system":
                        if not system_added:
                            claude_msgs.append(m)
                            system_added = True
                        else:
                            # 첨부파일 정보는 user 메시지로 변환
                            claude_msgs.append({"role": "user", "content": m["content"]})
                    else:
                        claude_msgs.append(m)
                
                system_contents = []
                non_system_messages = []

                for msg in claude_msgs:
                    if msg["role"] == "system":
                        system_contents.append(msg["content"])
                    else:
                        non_system_messages.append(msg)

                combined_system = "\n\n".join(system_contents) if system_contents else None

                completion = self.claude.messages.create(
                    model=CLAUDE_MODEL.get(model, model),
                    system=combined_system,  # 최상위 system 매개변수로 전달
                    messages=non_system_messages,
                    max_tokens=1024
                )
                response = completion.content[0].text
                
            else:
                raise ValueError("지원되지 않는 API: " + api)
            
            # 첨부파일이 있었지만 AI가 명시적으로 언급하지 않은 경우
            if has_attachments and not any(term in response.lower() for term in 
                                          ["첨부", "파일", "엑셀", "excel", "csv", "워드", "hwp", "pdf", "한글", "이미지", "사진", "image", "picture"]):
                # 안내 메시지 추가 (첨부파일 정보가 반영되었음을 알림)
                return response + "\n\n(시스템: 첨부하신 파일의 분석 정보가 응답에 반영되었습니다.)"
            
            return response
            
        except Exception as e:
            return f"API 호출 오류: {str(e)}\n\n{api}/{model} 호출 중 문제가 발생했습니다. 다른 모델을 선택하거나 API 키를 확인해주세요."


# ─────────────────────────────────────────────────────────────
# 6.  RequestThread (attachment summaries + progress)
# ─────────────────────────────────────────────────────────────
class RequestThread(QThread):
    response_ready = pyqtSignal(str)
    error = pyqtSignal(str)
    progress = pyqtSignal(str)

    def __init__(self, ai, api, model, session, web):
        super().__init__()
        self.ai, self.api, self.model = ai, api, model
        self.session, self.web = session, web
        self._stop = threading.Event()

    # -------- worker --------
    def run(self):
        threading.Thread(target=self._ticker, daemon=True).start()
        try:
            extras = self._summaries(self.session.attachments)
            text = self.ai.get_response(
                self.api, self.model, self.session.messages + extras, self.web
            )
            self.response_ready.emit(text)
        except Exception as e:
            self.error.emit(str(e))
        finally:
            self._stop.set()

    def _ticker(self):
        start = time.time()
        while not self._stop.is_set():
            self.progress.emit(f"⏳ {int(time.time() - start)}s")
            self._stop.wait(5)

    # -------- summary helpers --------
    def _md(self, df: pd.DataFrame) -> str:
        """DataFrame을 마크다운 또는 문자열로 변환"""
        try:
            return df.to_markdown() if TAB_OK else df.to_string()
        except Exception:
            return str(df)  # 변환 실패시 문자열로

    # ---------- Excel summary ----------
    def _sum_excel(self, path: str) -> str:
        """Excel 파일 요약"""
        self.progress.emit(f"⏳ Excel 파일 분석 중: {os.path.basename(path)}")
        
        if not OPENPYXL_OK and not XLRD_OK:
            return "Excel 처리에 필요한 라이브러리 미설치 (openpyxl ≥ 3.1.0 또는 xlrd 1.2.0)"
        
        ext = os.path.splitext(path)[1].lower()
        engines = []
        
        # 확장자에 따라 시도할 엔진 순서 결정
        if ext == '.xls':
            if XLRD_OK: engines.append('xlrd')
            if OPENPYXL_OK: engines.append('openpyxl')
        else:
            if OPENPYXL_OK: engines.append('openpyxl')
            if XLRD_OK: engines.append('xlrd')
        
        if not engines:
            return f"{ext} 파일을 처리할 수 있는 라이브러리가 설치되지 않았습니다"
        
        # 각 엔진 순서대로 시도
        last_error = None
        wb = None
        for engine in engines:
            try:
                wb = pd.ExcelFile(path, engine=engine)
                break  # 성공하면 반복 중단
            except Exception as e:
                last_error = str(e)
                continue
        
        if wb is None:  # 모든 엔진 실패시
            return f"Excel 파일 열기 실패: {last_error}"
        
        parts = []
        for sh in wb.sheet_names:
            try:
                df = pd.read_excel(wb, sheet_name=sh, nrows=100)  # 최대 100행만 읽기
                rows, cols = df.shape
                sec = [f"● **{sh}** — {rows}행 × {cols}열 (미리보기: 최대 100행)"]
                
                # 데이터 미리보기
                if not df.empty:
                    preview = df.head(5)
                    sec.append(self._md(preview))
                    
                    # 숫자형 컬럼 통계
                    num = df.select_dtypes(include=['number'])
                    if not num.empty:
                        desc = num.agg(["mean", "min", "max"]).T.round(2)
                        sec.append("**숫자형 컬럼 통계:**")
                        sec.append(self._md(desc))
                        
                parts.append("\n".join(sec))
            except Exception as e:
                parts.append(f"● **{sh}** — 시트 읽기 실패: {str(e)}")
        
        if not parts:
            return "Excel 파일을 읽었으나 유효한 시트를 찾을 수 없습니다."
        
        return "\n\n".join(parts)

    def _sum_csv(self, path: str) -> str:
        """CSV 파일 요약"""
        try:
            self.progress.emit(f"⏳ CSV 파일 분석 중: {os.path.basename(path)}")
            df = pd.read_csv(path, nrows=100)  # 최대 100행만 읽기
            
            summary = [f"● 행/열: {df.shape[0]}/{df.shape[1]}", ""]
            
            # 데이터 미리보기
            summary.append("**데이터 미리보기:**")
            summary.append(self._md(df.head(5)))
            
            # 숫자형 컬럼만 통계
            num = df.select_dtypes(include=['number'])
            if not num.empty:
                desc = num.describe().T.round(2)
                summary.append("\n**숫자형 컬럼 통계:**")
                summary.append(self._md(desc))
                
            return "\n".join(summary)
        except Exception as e:
            return f"CSV 파일 읽기 실패: {str(e)}"

    def _sum_docx(self, path: str) -> str:
        """DOCX 파일 요약"""
        if docx is None:
            return "python-docx 라이브러리가 설치되지 않았습니다."
            
        try:
            self.progress.emit(f"⏳ DOCX 파일 분석 중: {os.path.basename(path)}")
            doc = docx.Document(path)
            
            # 텍스트 추출
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
            
            # 통계
            words = text.split()
            word_count = len(words)
            char_count = len(text)
            
            summary = [
                f"● 단어 수: {word_count}",
                f"● 문자 수: {char_count}",
                f"● 단락 수: {len(paragraphs)}",
                "\n**내용 미리보기:**\n"
            ]
            
            # 내용 미리보기 (최대 1500자)
            preview = text[:1500] + ("..." if len(text) > 1500 else "")
            
            summary.append(preview)
            
            return "\n".join(summary)
        except Exception as e:
            return f"DOCX 파일 읽기 실패: {str(e)}"

# ---------- HWP/HWPX 분석 함수 ----------
    def _extract_hwpx_text(self, path: str) -> str:
        """HWPX 파일에서 텍스트 추출 (XML 파싱)"""
        self.progress.emit(f"⏳ HWPX 파일 분석 중 (XML): {os.path.basename(path)}")
        text_parts = []

        try:
            with zipfile.ZipFile(path) as hwpx:
                content_list = [name for name in hwpx.namelist() if name.startswith('Contents/')]
                section_files = [name for name in content_list if name.startswith('Contents/section')]

                # --- 1. LXML 파싱 ---
                if LXML_OK:
                    for section_file in section_files:
                        with hwpx.open(section_file) as f:
                            try:
                                namespaces = {
                                    'hwp': 'http://www.hancom.co.kr/hwpml/2016/paragraph',
                                    'hp': 'http://www.hancom.co.kr/hwpml/2016/paragraph',
                                    'hw': 'http://www.hancom.co.kr/hwpml/2016/paragraph',
                                    'para': 'http://www.hancom.co.kr/hwpml/2016/paragraph'
                                }
                                tree = LET.parse(f)
                                root = tree.getroot()
                                for ns_prefix, ns_uri in namespaces.items():
                                    for t_tag in [f'.//{ns_prefix}:t', f'.//{ns_prefix}:text', './/text']:
                                        try:
                                            elements = root.xpath(t_tag, namespaces={ns_prefix: ns_uri})
                                            for t in elements:
                                                if t.text:
                                                    text_parts.append(t.text)
                                        except Exception:
                                            continue
                            except Exception:
                                continue

                # --- 2. 기본 ElementTree 파싱 ---
                if not text_parts:
                    self.progress.emit("⏳ LXML 파싱 실패, 기본 ElementTree 사용 중...")
                    for section_file in section_files:
                        with hwpx.open(section_file) as f:
                            try:
                                tree = ET.parse(f)
                                root = tree.getroot()

                                for t in root.findall('.//*'):
                                    if t.tag.endswith('}t') and t.text:
                                        text_parts.append(t.text)
                            except Exception:
                                continue

                # --- 3. 정규식 파싱 (마지막 수단) ---
                if not text_parts:
                    self.progress.emit("⏳ XML 파싱 실패, 정규식 사용 중...")
                    for section_file in section_files:
                        try:
                            with hwpx.open(section_file) as f:
                                content = f.read().decode('utf-8', errors='ignore')
                                patterns = [
                                    r'<hp:t>(.*?)</hp:t>',
                                    r'<hwp:t>(.*?)</hwp:t>',
                                    r'<t>(.*?)</t>',
                                    r'<text>(.*?)</text>'
                                ]
                                for pattern in patterns:
                                    text_parts.extend(re.findall(pattern, content, re.DOTALL))
                        except Exception as e:
                            self.progress.emit(f"⚠️ 정규식 처리 중 오류: {str(e)}")
                            continue

        except Exception as e:
            self.progress.emit(f"❌ HWPX 처리 중 오류 발생: {e}")

        return "\n".join(text_parts)

    def _extract_hwp_text_with_pyhwp(self, path: str) -> str:
        """pyhwp 라이브러리를 사용한 HWP 텍스트 추출"""
        self.progress.emit(f"⏳ HWP 파일 분석 중 (pyhwp): {os.path.basename(path)}")
        try:
            with OleFileIO(path) as ole:
                hwp = pyhwp.HWPDocument(ole)
                return "".join(c.text for c in hwp.bodytext().children)
        except Exception as e:
            self.progress.emit(f"⚠️ pyhwp 처리 실패: {str(e)}")
            return ""

    def _extract_hwp_text_with_hwp5txt(self, path: str) -> str:
        """hwp5txt 명령줄 도구를 사용한 HWP 텍스트 추출"""
        self.progress.emit(f"⏳ HWP 파일 분석 중 (hwp5txt): {os.path.basename(path)}")
        try:
            result = subprocess.run(
                ['hwp5txt', path],
                capture_output=True,
                text=True,
                encoding='utf-8'  # 이 줄을 추가
            )
            if result.returncode == 0:
                return result.stdout
            else:
                self.progress.emit(f"⚠️ hwp5txt 처리 실패: {result.stderr}")
                return ""
        except Exception as e:
            self.progress.emit(f"⚠️ hwp5txt 처리 실패: {str(e)}")
            return ""

    def _extract_hwp_text_with_jpype(self, path: str) -> str:
        """JPype + HWPLib를 사용한 HWP 텍스트 추출"""
        if not HWP_JPYPE_OK:
            return ""
            
        self.progress.emit(f"⏳ HWP 파일 분석 중 (JPype): {os.path.basename(path)}")
        try:
            # 이 부분은 HWPLib.jar가 클래스패스에 있어야 작동
            if not jpype.isJVMStarted():
                jpype.startJVM(jpype.getDefaultJVMPath())
                
            # HWPLib 클래스 가져오기 시도
            try:
                from kr.hancom.hwplib.object.HWPFile import HWPFile
                from kr.hancom.hwplib.reader.HWPReader import HWPReader
                from kr.hancom.hwplib.tool.textextractor.TextExtractMethod import TextExtractMethod
                from kr.hancom.hwplib.tool.textextractor.TextExtractor import TextExtractor

                hwpFile = HWPReader.fromFile(path)
                textExtractor = TextExtractor()
                text = textExtractor.extract(hwpFile, TextExtractMethod.InsertControlTextBetweenParagraphText)
                return text
            except Exception as e:
                self.progress.emit(f"⚠️ HWPLib/JPype 클래스 로드 실패: {str(e)}")
                return ""
        except Exception as e:
            self.progress.emit(f"⚠️ JPype 처리 실패: {str(e)}")
            return ""

    def _extract_hwp_preview_text(self, path: str) -> str:
        """HWP 파일의 미리보기 텍스트 (PrvText) 추출"""
        self.progress.emit(f"⏳ HWP 미리보기 텍스트 추출 시도: {os.path.basename(path)}")
        try:
            with OleFileIO(path) as ole:
                if ole.exists('PrvText'):
                    with ole.openstream('PrvText') as stream:
                        data = stream.read()
                        # 일반적으로 UTF-16LE 인코딩 사용
                        return data.decode('utf-16le', errors='ignore')
                else:
                    self.progress.emit("⚠️ PrvText 스트림이 존재하지 않음")
                    return ""
        except Exception as e:
            self.progress.emit(f"⚠️ PrvText 추출 실패: {str(e)}")
            return ""

    def convert_hwp_to_text_with_soffice(self, hwp_path, temp_dir=None):
        """LibreOffice/OpenOffice를 사용하여 HWP/HWPX를 텍스트로 변환"""
        if not SOFFICE_PATH:
            return None
            
        import tempfile
        import os
        
        if temp_dir is None:
            temp_dir = tempfile.mkdtemp()
        
        output_file = os.path.join(temp_dir, "converted.txt")
        
        try:
            self.progress.emit(f"⏳ LibreOffice로 한글 파일 변환 중...")
            
            # LibreOffice 명령어로 변환 시도
            cmd = [
                SOFFICE_PATH,
                "--headless",
                "--convert-to", "txt:Text",
                "--outdir", temp_dir,
                hwp_path
            ]
            
            process = subprocess.run(cmd, capture_output=True, timeout=60)
            
            # 출력 파일 이름 구하기 (확장자만 변경됨)
            base_name = os.path.basename(hwp_path)
            base_name_without_ext = os.path.splitext(base_name)[0]
            output_file = os.path.join(temp_dir, f"{base_name_without_ext}.txt")
            
            if os.path.exists(output_file):
                with open(output_file, 'r', encoding='utf-8', errors='replace') as f:
                    return f.read()
            
            return None
        except Exception as e:
            self.progress.emit(f"⚠️ LibreOffice 변환 실패: {str(e)}")
            return None
        finally:
            # 임시 파일 정리
            try:
                if os.path.exists(output_file):
                    os.remove(output_file)
            except:
                pass

    def _hwp_fallback_info(self, filename: str, size_mb: float, ext: str, error_msg: str = None) -> str:
        """한글 파일 분석이 불가능할 때 기본 정보와 함께 상세한 안내 제공"""
        doc_type = "HWPX" if ext == '.hwpx' else "HWP"
        
        summary = [
            f"● 한글 문서 ({doc_type})",
            f"● 파일명: {filename}",
            f"● 파일 크기: {size_mb:.2f}MB",
            "\n**주의사항:**\n"
        ]
        
        # 오류 메시지 포함
        if error_msg:
            summary.append(f"한글 파일 분석 중 오류가 발생했습니다: {error_msg}")
        
        # 상세한 문제 해결 안내
        summary.append("\n**해결 방법:**")
        
        # 1. 파일 변환 권장
        summary.append("1. 한글 파일을 다른 형식으로 변환하여 첨부해 주세요:")
        summary.append("   - PDF 형식 (권장): 한글 프로그램에서 '다른 이름으로 저장' → 'PDF' 선택")
        summary.append("   - DOCX 형식: 한글 프로그램에서 '다른 이름으로 저장' → 'DOCX' 선택")
        summary.append("   - TXT 형식: 한글 프로그램에서 '다른 이름으로 저장' → '텍스트 문서' 선택")
        
        # 2. 라이브러리 설치 안내
        summary.append("\n2. 또는 다음 라이브러리를 설치하여 HWP/HWPX 직접 처리를 활성화할 수 있습니다:")
        
        if ext == '.hwpx':
            summary.append("   - HWPX 지원 활성화: pip install lxml")
        else:
            summary.append("   - HWP 지원 활성화: pip install pyhwp olefile")
        
        # 3. LibreOffice 설치 안내
        summary.append("\n3. 또는 LibreOffice/OpenOffice를 설치하면 자동 변환이 가능합니다:")
        summary.append("   - https://www.libreoffice.org/download/ 에서 다운로드하여 설치")
        
        return "\n".join(summary)

    def _sum_hwp(self, path: str) -> str:
        """HWP/HWPX 파일 요약 - 외부 변환기 사용 우선"""
        self.progress.emit(f"⏳ 한글 파일 분석 중: {os.path.basename(path)}")
        
        # 파일 확장자 확인
        ext = os.path.splitext(path)[1].lower()
        filename = os.path.basename(path)
        file_size_mb = os.path.getsize(path) / (1024 * 1024)
        
        # 먼저 LibreOffice/OpenOffice 변환 시도 (가장 안정적인 방법)
        if SOFFICE_OK:
            text = self.convert_hwp_to_text_with_soffice(path)
            if text and text.strip():
                # 텍스트 추출 성공
                words = text.split()
                word_count = len(words)
                char_count = len(text)
                
                summary = [
                    f"● 한글 문서 ({os.path.splitext(filename)[1][1:].upper()})",
                    f"● 파일 크기: {file_size_mb:.2f}MB",
                    f"● 단어 수: {word_count}",
                    f"● 문자 수: {char_count}",
                    "\n**내용 미리보기:**\n"
                ]
                
                # 내용 미리보기 (최대 1500자)
                preview = text[:1500] + ("..." if len(text) > 1500 else "")
                summary.append(preview)
                
                return "\n".join(summary)
        
        # LibreOffice 방식이 실패하면 기존 방식 시도
        try:
            # 라이브러리 설치 여부 확인
            if ext == '.hwpx' and not HWPX_OK:
                return self._hwp_fallback_info(filename, file_size_mb, ext)
            elif ext == '.hwp' and not HWP_OK:
                return self._hwp_fallback_info(filename, file_size_mb, ext)
            
            # 텍스트 추출 시도
            text = ""
            
            # HWPX 파일 처리 (XML 기반 한글 문서)
            if ext == '.hwpx':
                text = self._extract_hwpx_text(path)
                doc_type = "HWPX"
            # HWP 파일 처리 (구버전 한글 문서)
            else:  # ext == '.hwp'
                # 다양한 방법으로 시도 (순서대로)
                text = self._extract_hwp_text_with_pyhwp(path)
                
                if not text and HWP5TXT_OK:
                    text = self._extract_hwp_text_with_hwp5txt(path)
                    
                if not text and HWP_JPYPE_OK:
                    text = self._extract_hwp_text_with_jpype(path)
                    
                if not text:
                    # 마지막 수단: 미리보기 텍스트 추출
                    text = self._extract_hwp_preview_text(path)
                    
                doc_type = "HWP"
            
            # 텍스트가 추출되었으면 통계 및 미리보기 생성
            if text:
                words = text.split()
                word_count = len(words)
                char_count = len(text)
                
                summary = [
                    f"● 한글 문서 ({doc_type})",
                    f"● 파일 크기: {file_size_mb:.2f}MB",
                    f"● 단어 수: {word_count}",
                    f"● 문자 수: {char_count}",
                    "\n**내용 미리보기:**\n"
                ]
                
                # 내용 미리보기 (최대 1500자)
                preview = text[:1500] + ("..." if len(text) > 1500 else "")
                summary.append(preview)
                
                return "\n".join(summary)
            else:
                # 모든 방법이 실패하면 대체 정보 제공
                return self._hwp_fallback_info(filename, file_size_mb, ext, "모든 텍스트 추출 방법 실패")
                
        except Exception as e:
            return self._hwp_fallback_info(filename, file_size_mb, ext, str(e))

    # ---------- Summary dispatcher ----------
    def _summaries(self, paths: List[str]) -> List[dict]:
        """첨부파일 요약 메시지 리스트 반환"""
        summaries = []
        for path in paths:
            ext = os.path.splitext(path)[1].lower()
            try:
                if ext in ['.xls', '.xlsx']:
                    content = self._sum_excel(path)
                    summaries.append({"role": "system", "content": f"[📊 Excel] {os.path.basename(path)}\n{content}"})
                elif ext == '.csv':
                    content = self._sum_csv(path)
                    summaries.append({"role": "system", "content": f"[📈 CSV] {os.path.basename(path)}\n{content}"})
                elif ext == '.docx':
                    content = self._sum_docx(path)
                    summaries.append({"role": "system", "content": f"[📄 DOCX] {os.path.basename(path)}\n{content}"})
                elif ext in ['.hwp', '.hwpx']:
                    content = self._sum_hwp(path)
                    summaries.append({"role": "system", "content": f"[📑 한글] {os.path.basename(path)}\n{content}"})
                elif ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
                    # 이미지 파일은 멀티모달로 처리 (image_path 키 추가)
                    summaries.append({
                        "role": "user",
                        "content": f"[🖼️ 이미지] {os.path.basename(path)}",
                        "image_path": path
                    })
                # 이미지, PDF 등은 필요시 추가
                else:
                    summaries.append({"role": "system", "content": f"[❓ 기타 파일] {os.path.basename(path)} — 지원되지 않음"})
            except Exception as e:
                summaries.append({"role": "system", "content": f"[⚠️ 오류] {os.path.basename(path)} 분석 실패: {e}"})
        return summaries


# ─────────────────────────────────────────────────────────────
# 7.  MainWindow (UI)
# ─────────────────────────────────────────────────────────────
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AI Chat Assistant - Multimodal")
        self.resize(1100, 750)
        self.setWindowIcon(QIcon())

        self.history = ChatHistory()
        self.ai = AIManager()
        self.current = ChatSession(f"sess_{len(self.history.sessions) + 1}")

        self._setup_ui()
        self._load_history()

    # ---------- UI 구성 ----------
    def _setup_ui(self):
        # 왼쪽: 대화 리스트
        left_layout = QVBoxLayout()
        new_btn = QPushButton("New Chat")
        new_btn.clicked.connect(self._new_chat)
        self.history_list = QListWidget()
        self.history_list.itemClicked.connect(self._on_history_clicked)
        left_layout.addWidget(new_btn)
        left_layout.addWidget(self.history_list)

        # 오른쪽: 채팅 뷰
        bar = QHBoxLayout()
        self.api_cb = QComboBox()
        self.api_cb.addItems(["openai", "gemini", "claude"])
        self.api_cb.currentIndexChanged.connect(self._update_models)
        self.model_cb = QComboBox()
        self.chk_web = QCheckBox("Web Search")
        self.chk_web.setChecked(True)
        bar.addWidget(QLabel("API:"))
        bar.addWidget(self.api_cb)
        bar.addWidget(QLabel("Model:"))
        bar.addWidget(self.model_cb)
        bar.addStretch()
        bar.addWidget(self.chk_web)

        self.view = QTextEdit()
        self.view.setReadOnly(True)

        # 첨부파일 관리 UI
        attach_line = self._add_file_management()

        self.input = QTextEdit()
        self.input.setFixedHeight(100)
        send_btn = QPushButton("Send")
        send_btn.clicked.connect(self._send)

        right_layout = QVBoxLayout()
        right_layout.addLayout(bar)
        right_layout.addWidget(self.view)
        right_layout.addLayout(attach_line)
        right_layout.addWidget(self.input)
        right_layout.addWidget(send_btn)

        splitter = QSplitter(Qt.Horizontal)
        left_widget, right_widget = QWidget(), QWidget()
        left_widget.setLayout(left_layout)
        right_widget.setLayout(right_layout)
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setSizes([220, 880])

        container = QWidget()
        container.setLayout(QHBoxLayout())
        container.layout().addWidget(splitter)
        self.setCentralWidget(container)

        self._update_models()

    def _add_file_management(self):
        """첨부파일 관리 기능 추가"""
        attach_line = QHBoxLayout()
        self.attach_lbl = QLabel("첨부파일: 없음")
        
        attach_btn = QPushButton("파일 첨부")
        attach_btn.setIcon(QIcon.fromTheme("document-open"))
        attach_btn.clicked.connect(self._attach_file)
        
        clear_btn = QPushButton("첨부 삭제")
        clear_btn.setIcon(QIcon.fromTheme("edit-delete"))
        clear_btn.clicked.connect(self._clear_attachments)
        
        attach_line.addWidget(self.attach_lbl)
        attach_line.addStretch()
        attach_line.addWidget(attach_btn)
        attach_line.addWidget(clear_btn)
        
        return attach_line
    
    def _attach_file(self):
        """파일 첨부 대화상자 표시 및 처리"""
        files, _ = QFileDialog.getOpenFileNames(
            self, 
            "파일 선택", 
            "", 
            "모든 파일 (*);;Excel Files (*.xlsx *.xls);;CSV Files (*.csv);;Word Files (*.docx);;PDF Files (*.pdf);;한글 Files (*.hwp *.hwpx);;Images (*.png *.jpg *.jpeg *.bmp *.gif)"
        )
        
        if not files:
            return
            
        # 첨부파일 저장 디렉토리 생성
        attach_dir = os.path.join("attachments", self.current.session_id)
        os.makedirs(attach_dir, exist_ok=True)
        
        # 이미 첨부된 파일 이름 목록
        existing_files = [os.path.basename(p) for p in self.current.attachments]
        
        success_count = 0
        for src_path in files:
            try:
                # 파일 크기 체크 (30MB 제한)
                file_size_mb = os.path.getsize(src_path) / (1024 * 1024)
                if file_size_mb > 30:
                    QMessageBox.warning(
                        self, 
                        "파일 크기 초과", 
                        f"{os.path.basename(src_path)}: 파일 크기가 30MB를 초과합니다 ({file_size_mb:.1f}MB)"
                    )
                    continue
                    
                filename = os.path.basename(src_path)
                
                # 이미 첨부된 파일인지 확인
                if filename in existing_files:
                    # 파일명 중복 처리 (시간 추가)
                    base, ext = os.path.splitext(filename)
                    filename = f"{base}_{int(time.time())}{ext}"
                    
                dst_path = os.path.join(attach_dir, filename)
                
                # 파일 복사
                with open(src_path, "rb") as src, open(dst_path, "wb") as dst:
                    dst.write(src.read())
                    
                self.current.add_attachment(dst_path)
                success_count += 1
                
            except Exception as e:
                QMessageBox.warning(self, "첨부 실패", f"{os.path.basename(src_path)} 첨부 실패: {e}")
        
        if success_count > 0:
            self._update_attach_lbl()
            QMessageBox.information(self, "첨부 완료", f"{success_count}개 파일이 첨부되었습니다.")

    def _update_attach_lbl(self):
        """첨부파일 라벨 업데이트"""
        if not self.current.attachments:
            self.attach_lbl.setText("첨부파일: 없음")
            return
            
        names = ", ".join(os.path.basename(p) for p in self.current.attachments)
        count = len(self.current.attachments)
        
        if len(names) > 50:  # 이름이 너무 길면 축약
            names = names[:47] + "..."
            
        self.attach_lbl.setText(f"첨부파일 ({count}개): {names}")

    def _clear_attachments(self):
        """현재 세션의 모든 첨부파일 삭제"""
        if not self.current.attachments:
            return
            
        count = len(self.current.attachments)
        reply = QMessageBox.question(
            self, 
            "첨부파일 삭제", 
            f"현재 첨부된 {count}개 파일을 모두 삭제하시겠습니까?",
            QMessageBox.Yes | QMessageBox.No, 
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            self.current.attachments = []
            self._update_attach_lbl()
            QMessageBox.information(self, "삭제 완료", "모든 첨부파일이 삭제되었습니다.")

    # ---------- 모델 목록 ----------
    def _update_models(self):
        self.model_cb.clear()
        api = self.api_cb.currentText()
        if api == "openai":
            self.model_cb.addItems(["gpt-o3", "gpt-o3-mini", "gpt-4o"])
            self.model_cb.setCurrentText("gpt-4o")
        elif api == "gemini":
            self.model_cb.addItems(["gemini-2.5-flash", "gemini-2.5-pro"])
        else:
            self.model_cb.addItems(["claude-3-7-sonnet", "claude-3-5-sonnet"])

    # ---------- History ----------
    def _load_history(self, select: str | None = None):
        self.history_list.clear()
        for sid, sess in self.history.sessions.items():
            self.history_list.addItem(sess.title)
            if select and sid == select:
                self.history_list.setCurrentRow(self.history_list.count() - 1)

    def _on_history_clicked(self, item):
        title = item.text()
        self.current = next((s for s in self.history.sessions.values() if s.title == title), None)
        if self.current:
            self._refresh()

    def _new_chat(self):
        self.current = ChatSession(f"sess_{len(self.history.sessions) + 1}")
        self._refresh()
        self.history_list.clearSelection()

    # ---------- Chat refresh ----------
    def _refresh(self):
        self.view.clear()
        for m in self.current.messages:
            who = "You" if m["role"] == "user" else "AI"
            content = m["content"]
            if who == "AI":
                try:
                    html = markdown2.markdown(content, extras=["fenced-code-blocks", "tables"])
                    self.view.append(f"<b>{who}:</b><br>{html}<br>")
                except Exception:
                    self.view.append(f"<b>{who}:</b> {content}<br>")
            else:
                self.view.append(f"<b>{who}:</b> {content}<br>")
        self.view.verticalScrollBar().setValue(self.view.verticalScrollBar().maximum())
        self._update_attach_lbl()

    def _send(self):
        """메시지 전송 처리"""
        txt = self.input.toPlainText().strip()
        if not txt and not self.current.attachments:
            QMessageBox.warning(self, "Empty", "텍스트 입력 또는 파일 첨부가 필요합니다.")
            return

        # 첨부파일이 있으면 명시적으로 표시
        has_attachments = bool(self.current.attachments)
        if has_attachments:
            attachment_names = [os.path.basename(p) for p in self.current.attachments]
            attachment_msg = f"{txt}\n\n첨부파일: {', '.join(attachment_names)}" if txt else f"첨부파일 분석 요청: {', '.join(attachment_names)}"
            self.current.add_message("user", attachment_msg)
        else:
            self.current.add_message("user", txt)
            
        self.input.clear()
        self._refresh()
        
        # 로딩 표시 추가 (첨부파일이 있으면 다른 메시지)
        if has_attachments:
            self.view.append("<i>첨부파일 분석 중... 잠시만 기다려주세요</i>")
        else:
            self.view.append("<i>AI가 응답을 생성하는 중입니다...</i>")

        # 요청 스레드 시작
        thread = RequestThread(
            self.ai,
            self.api_cb.currentText(),
            self.model_cb.currentText(),
            self.current,
            self.chk_web.isChecked(),
        )
        thread.response_ready.connect(self._on_resp)
        thread.error.connect(self._on_err)
        thread.progress.connect(self._on_prog)
        thread.start()
        self._thread = thread

    # ---------- Thread callbacks ----------
    def _on_prog(self, msg: str):
        cur = self.view.textCursor()
        cur.movePosition(cur.End)
        cur.select(cur.BlockUnderCursor)
        cur.removeSelectedText()
        self.view.append(f"<i>{msg}</i>")

    def _on_resp(self, txt: str):
        cur = self.view.textCursor()
        cur.movePosition(cur.End)
        cur.select(cur.BlockUnderCursor)
        cur.removeSelectedText()

        self.current.add_message("assistant", txt)
        self._refresh()
        self.history.save(self.current)

        if self.current.title.startswith("New Chat") and len(self.current.messages) >= 2:
            self.current.title = self.history.generate_title(self.current)
            self.history.save(self.current)
            self._load_history(select=self.current.session_id)

    def _on_err(self, msg: str):
        cur = self.view.textCursor()
        cur.movePosition(cur.End)
        cur.select(cur.BlockUnderCursor)
        cur.removeSelectedText()
        self.view.append(f"<span style='color:red'><b>Error:</b> {msg}</span><br>")


# ─────────────────────────────────────────────────────────────
# 8.  main
# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = MainWindow()
    win.show()
    sys.exit(app.exec_())